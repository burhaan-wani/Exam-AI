import base64
import io
import logging

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


def _data_url_from_bytes(file_bytes: bytes, mime_type: str) -> str:
    encoded = base64.b64encode(file_bytes).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def extract_images_from_pdf(file_bytes: bytes) -> list[str]:
    """Extract embedded page images from a PDF as data URLs when available."""
    reader = PdfReader(io.BytesIO(file_bytes))
    image_urls: list[str] = []

    def _normalized_filters(image_obj) -> list[str]:
        filters = image_obj.get("/Filter")
        if isinstance(filters, list):
            return [str(item) for item in filters]
        if filters:
            return [str(filters)]
        return []

    def _mime_type_from_image(image_obj, image_name: str) -> str:
        lowered_name = (image_name or "").lower()
        filters = _normalized_filters(image_obj)

        if lowered_name.endswith(".jpg") or lowered_name.endswith(".jpeg") or "/DCTDecode" in filters:
            return "image/jpeg"
        if lowered_name.endswith(".jp2") or "/JPXDecode" in filters:
            return "image/jp2"
        if lowered_name.endswith(".webp"):
            return "image/webp"

        if lowered_name.endswith(".png") or "/FlateDecode" in filters:
            return "image/png"
        return "image/jpeg"

    def _append_image(image_obj, image_name: str) -> None:
        mime_type = _mime_type_from_image(image_obj, image_name)

        data = None
        raw_data = getattr(image_obj, "_data", None)
        if raw_data and mime_type in {"image/jpeg", "image/jp2", "image/webp"}:
            data = raw_data
        if not data:
            try:
                data = image_obj.get_data()
            except Exception:
                data = None
        if not data:
            return

        encoded = base64.b64encode(data).decode("ascii")
        image_urls.append(f"data:{mime_type};base64,{encoded}")

    def _extract_from_xobjects(xobjects) -> None:
        try:
            resolved = xobjects.get_object()
        except Exception:
            resolved = xobjects

        for image_name, image_obj in (resolved or {}).items():
            try:
                image_obj = image_obj.get_object()
            except Exception:
                pass

            subtype = str(image_obj.get("/Subtype", ""))
            if subtype == "/Image":
                _append_image(image_obj, str(image_name))
                continue

            if subtype == "/Form":
                nested_resources = image_obj.get("/Resources")
                if nested_resources:
                    nested_xobjects = nested_resources.get("/XObject")
                    if nested_xobjects:
                        _extract_from_xobjects(nested_xobjects)

    for page in reader.pages:
        page_found_image = False
        try:
            images = getattr(page, "images", [])
        except Exception:
            images = []

        for image in images or []:
            data = getattr(image, "data", None)
            name = (getattr(image, "name", "") or "").lower()
            if not data:
                continue

            if name.endswith(".png"):
                mime_type = "image/png"
            elif name.endswith(".webp"):
                mime_type = "image/webp"
            else:
                mime_type = "image/jpeg"

            encoded = base64.b64encode(data).decode("ascii")
            image_urls.append(f"data:{mime_type};base64,{encoded}")
            page_found_image = True

        if page_found_image:
            continue

        try:
            resources = page.get("/Resources")
        except Exception:
            resources = None

        if not resources:
            continue

        xobjects = resources.get("/XObject")
        if xobjects:
            _extract_from_xobjects(xobjects)

    return image_urls


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    parts.extend(p.text.strip() for p in doc.paragraphs if p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_lines = [line.strip() for line in cell.text.splitlines() if line.strip()]
                parts.extend(cell_lines)

    return "\n\n".join(parts)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from an uploaded file based on extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    logger.info(f"Extracting text from '{filename}' (type={ext})")

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Use PDF, DOCX, or TXT.")


def extract_template_source(file_bytes: bytes, filename: str) -> tuple[str, list[str]]:
    """
    Extract text and any embedded page images for template parsing.
    Image extraction is primarily useful for scanned PDFs that contain little
    or no selectable text.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in {"png", "jpg", "jpeg", "webp"}:
        mime_type = {
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "webp": "image/webp",
        }[ext]
        return "", [_data_url_from_bytes(file_bytes, mime_type)]
    if ext != "pdf":
        return extract_text(file_bytes, filename), []
    return extract_text_from_pdf(file_bytes), extract_images_from_pdf(file_bytes)
